import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_markdown(text):
    # ลบตัวหนา/เอียง (***, **, *, __, _)
    text = re.sub(r'\*+', '', text)
    text = re.sub(r'_+', '', text)
    # ลบ inline code (`)
    text = re.sub(r'`+', '', text)
    return text.strip()

def create_word_from_md(md_file, output_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    doc = Document()
    
    # Set default font to something standard
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # จัดการรูปภาพ (รูปแบบ: ![caption](path/to/image))
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=doc.sections[0].page_width * 0.7) # ขนาด 70% ของหน้า
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # ใส่ Caption ใต้รูป
                cap_p = doc.add_paragraph(f"รูปที่ {caption}")
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # จัดการ Heading
        if line.startswith("# "):
            doc.add_heading(clean_markdown(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(clean_markdown(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_markdown(line[4:]), level=2)
        # จัดการ List
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(clean_markdown(line[2:]), style='List Bullet')
        # ข้อความทั่วไป
        else:
            doc.add_paragraph(clean_markdown(line))

    doc.save(output_file)
    print(f"Created Word document: {output_file}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python word_generator.py <input_md> <output_docx>")
    else:
        create_word_from_md(sys.argv[1], sys.argv[2])
